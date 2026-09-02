import argparse,json
from pathlib import Path
from docx import Document

parser=argparse.ArgumentParser()
parser.add_argument('docx')
parser.add_argument('--out',required=True)
args=parser.parse_args()
doc=Document(args.docx)
paragraphs=[{'index':i+1,'style':p.style.name,'text':p.text.strip()} for i,p in enumerate(doc.paragraphs) if p.text.strip()]
tables=[]
for ti,table in enumerate(doc.tables,1):
    rows=[]
    for row in table.rows:
        rows.append(['\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells])
    tables.append({'index':ti,'rows':rows})
payload={'source':str(Path(args.docx).resolve()),'paragraphs':paragraphs,'tables':tables,'inline_shapes':len(doc.inline_shapes)}
Path(args.out).write_text(json.dumps(payload,ensure_ascii=False,indent=2),encoding='utf-8')
print(f"Extracted {len(paragraphs)} paragraphs, {len(tables)} tables and {len(doc.inline_shapes)} inline images.")
